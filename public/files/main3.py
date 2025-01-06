from docx import Document
import json

def generate_question_doc(data, output_path):
    """
    Generates a Word document from a list of questions and answers.

    :param data: List of dictionaries with 'question', 'choices', and 'answer'.
    :param output_path: File path to save the generated document.
    """
    # Initialize the Document
    doc = Document()

    # Populate the document with questions and answers
    for idx, item in enumerate(data, start=1):
        # Add question
        doc.add_paragraph(f"{idx}. {item['question']}")
        
        # Add choices with correct answer bolded
        for choice in item['choices']:
            para = doc.add_paragraph(style='List Bullet')
            if choice == item['answer']:
                para.add_run(choice).bold = True
            else:
                para.add_run(choice)

    # Save the document
    doc.save(output_path)

# Example data
with open("questions.json", "r") as f:
    data = json.load(f)

# Specify the output file path
output_path = "Questions_and_Answers.docx"

# Generate the document
generate_question_doc(data, output_path)

print(f"Document has been saved at {output_path}")
